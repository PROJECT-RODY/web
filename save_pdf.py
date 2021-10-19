from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

def create_pdf(title, sentences, image_name):
    

    pdf = FPDF()
    
    # Adding a page
    pdf.add_page()

    pdf.add_font('malgun', '', '/usr/share/fonts/truetype/malgun.ttf', uni=True) # 한글 폰트 추가
    # set style and size of font 
    # pdf.set_font("malgun", size = 30)
    
    pdf.image('./static/user_image/'+image_name+'.png', x = 60, )# 이미지 가로축 위치 설정
    # create a cell
    pdf.set_font("malgun", size = 30)
    pdf.cell(200, 10, txt = "", ln = 1, align = 'C')
    pdf.cell(200, 10, txt = title, ln = 2, align = 'C')
    pdf.cell(200, 10, txt = "", ln = 3, align = 'C')
    
    
    # add another cell
    pdf.set_font("malgun", size = 15)
    sentence = ''
    num_c = 0
    for i, s in enumerate(sentences.split(' ')):
        sentence += s + " "
        if i % 8 == 0:
            pdf.cell(200, 10, txt = sentence, ln = 4+num_c, align = 'C')
            sentence = ''
            num_c+=1
            
    

    # save the pdf
    save_path = "./static/pdf/"+image_name+".pdf"
    pdf.output(save_path)

    return save_path



def create_word(title, sentence, image_name):
    document = Document()
    document.add_heading('제목 : {}'.format(title), 0)
    document.add_picture('./user_image/'+image_name+'.png') # 이미지 삽입
    # 이미지 가운데 정렬 시적
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 이미지 가운데 정렬 끝
    document.add_paragraph('{}'.format(sentence))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.save("./pdf/"+image_name+".docx")
    

if __name__ == '__main__':
    # create_word("나무꾼", "나무꾼이 나무를 한다.", "single_s9")
    create_pdf("나무꾼", "나무꾼이 나무를 한다.", "single_s9")
    